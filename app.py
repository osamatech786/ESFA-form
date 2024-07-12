import streamlit as st
from streamlit_drawable_canvas import st_canvas
import re
from PIL import Image as PILImage
import shutil
from docx import Document
from docx.shared import Inches
from datetime import datetime, date

def app():
    st.set_page_config(
        page_title="Online Form",
        page_icon="📝",
        layout="wide",
        initial_sidebar_state="collapsed",
    )

    # Custom CSS to set a light background color
    st.markdown("""
        <style>
        body {
            background-color: #f0f0f0; /* Light grey background */
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    st.image('header.jpg', use_column_width=True)

    st.title('Welcome')
    st.subheader('Please fill out the following details:')

    # Form inputs
    partnership = st.text_input('Partnership')
    learner_name = st.text_input('Learner Name')
    qualification = st.selectbox('Qualification', [
        'High School Diploma', 'Bachelor\'s Degree', 'Master\'s Degree', 'PhD', 'Other'
    ])
    start_date = st.date_input(
        label="Aim Start Date",
        value=datetime(2000, 1, 1),  # Default date
        min_value=date(1900, 1, 1),  # Minimum selectable date
        max_value=date(2025, 12, 31),  # Maximum selectable date
        key="date_input_widget",  # Unique key for the widget
        help="Choose a date"  # Tooltip text
    )
    end_date = st.date_input(
        label="Expected Aim End Date",
        value=datetime(2000, 1, 1),  # Default date
        min_value=date(1900, 1, 1),  # Minimum selectable date
        max_value=date(2025, 12, 31),  # Maximum selectable date
        help="Choose a date"  # Tooltip text
    )

    st.header(
        "Start Paperwork - S2 Participant Assessment, Planning and Support and Qualification Start")

    # Define checklist items with their respective checkbox states
    checklist_items = [
        ("Original ILR/ILP", "original_ilr_ilp"),
        ("Original Eligibility Form", "original_eligibility_form"),
        ("Initial Assessment Outcomes Form (With Literacy & Numeracy scores)", "assessment_outcomes"),
        ("Contact Log for Start of Learning", "contact_log"),
        ("Copy of PLR", "copy_of_plr"),
        ("Contact Log for Start of Learning (Timesheet)", "contact_log_timesheet")
    ]

    for doc, key_base in checklist_items:
        st.subheader(doc)
        col1, col2 = st.columns([1, 1])
        with col1:
            st.checkbox('Enclosed & Complete (✓)', key=f"{key_base}_enclosed")
        with col2:
            st.checkbox('Checked as correct by Prevista (✓)', key=f"{key_base}_checked")

    st.subheader('Provider Signature')

    st.text("Signature:")
    provider_signature = st_canvas(
        fill_color="rgba(255, 255, 255, 1)",  
        stroke_width=5,
        stroke_color="rgb(0, 0, 0)",  # Black stroke color
        background_color="white",  # White background color
        width=400,
        height=150,
        drawing_mode="freedraw",
        key="canvas",
    )

    provider_date_signed = st.date_input(
    label="Date",
    value=datetime(2000, 1, 1),  # Default date
    min_value=date(1900, 1, 1),  # Minimum selectable date
    max_value=date(2025, 12, 31),  # Maximum selectable date
    help="Choose a date"  # Tooltip text
)

    st.header('Individual Learner/Participant Record and Plan')
    st.subheader('Please complete in BLOCK CAPITALS')

    # Subcontractor name or N/A if internal
    subcontractor_name = st.text_input('Subcontractor name or N/A if internal')

    # National Insurance Number
    ni_number = st.text_input('National Insurance Number')

    # Title and Name
    st.subheader('Title and Name')
    title = st.selectbox('Title', ['Mr', 'Mrs', 'Miss', 'Ms'])
    first_name = st.text_input('First Name')
    middle_name = st.text_input('Middle Name')
    family_name = st.text_input('Family Name')

    # Date of Birth and Current Age
    st.subheader('Date of Birth and Age')
    dob = st.date_input('Date of Birth (DD/MM/YYYY)')
    current_age = st.number_input('Current Age at Start of programme', min_value=0)

    # Address
    st.subheader('Address')
    house_name_street = st.text_input('House No./Name & Street')
    suburb_village = st.text_input('Suburb / Village')
    town_city = st.text_input('Town / City')
    county = st.text_input('County')
    country_of_domicile = st.text_input('Country of Domicile')
    current_postcode = st.text_input('Current Postcode')
    postcode_prior_enrolment = st.text_input('Post code prior to enrolment')

    # Contact Information
    st.subheader('Contact Information')
    email_address = st.text_input('Email Address')
    primary_telephone = st.text_input('Primary Telephone Number')
    secondary_telephone = st.text_input('Secondary Telephone Number')

    # Sex
    st.subheader('Sex')
    sex = st.selectbox('Sex (select)', ['M', 'F', 'Other'])
    if sex == 'Other':
        sex_other = st.text_input('If Other, please state')

    # Ethnicity Section
    st.header('Ethnicity')

    ethnicity_options = {
        'White': {
            'English/ Welsh/ Scottish/ N Irish/ British': '31',
            'Irish': '32',
            'Roma, Gypsy or Irish Traveller': '33',
            'Any other white background': '34'
        },
        'Mixed/ Multiple ethnic group': {
            'White and Black Caribbean': '35',
            'White and Black African': '36',
            'White and Asian': '37',
            'Any other mixed/ multiple ethnic background': '38'
        },
        'Asian/ Asian British': {
            'Indian': '39',
            'Pakistani': '40',
            'Bangladeshi': '41',
            'Chinese': '42',
            'Any other Asian background': '43'
        },
        'Black/ African/ Caribbean/ Black British': {
            'African': '44',
            'Caribbean': '45',
            'Any Other Black/ African/ Caribbean background': '46'
        },
        'Other Ethnic Group': {
            'Arab': '47',
            'Any other ethnic group': '98'
        }
    }

    ethnicity_category = st.selectbox('Select Ethnicity Category', list(ethnicity_options.keys()))
    ethnicity = st.selectbox('Select Ethnicity', list(ethnicity_options[ethnicity_category].keys()))
    ethnicity_code = ethnicity_options[ethnicity_category][ethnicity]
    st.write(f'Ethnicity Code: {ethnicity_code}')

    # Prior Attainment Section
    st.header('Prior Attainment')

    prior_attainment_options = {
        'Entry Level': '1',
        'Level 1': '2',
        'Level 2': '3',
        'Full Level 2': '4',
        'Level 3': '5',
        'Full Level 3': '6',
        'Level 4': '7',
        'Level 5': '8',
        'Level 6': '9',
        'Level 7 and above': '10',
        'No Qualifications': '99'
    }

    prior_attainment = st.selectbox('Select Prior Attainment Level', list(prior_attainment_options.keys()))
    prior_attainment_code = prior_attainment_options[prior_attainment]
    st.write(f'Prior Attainment Code: {prior_attainment_code}')

    # Emergency Contact Section
    st.header('Next of Kin/Emergency Contact')

    next_of_kin = st.text_input('Next of kin/Emergency contact')
    emergency_contact_phone = st.text_input('Emergency Contact Phone Number')


    # Household Situation Section
    st.header('Household Situation')
    st.subheader('Please select the most relevant option (place an "x" in ALL relevant boxes)')

    household_options = {
        '1 - No household member in employment with one or more dependent children': 'JH, JH+DC',
        '2 - No household member in employment with no dependent children': 'JH',
        '3 - Participant lives in a single adult household with dependent children': 'SAH+DC',
        '4 - Learner lives in single unemployed adult household with dependent children': 'JH, SAH+DC',
        '99 - None of the above apply': 'N/A'
    }

    household_selections = {}
    for option, code in household_options.items():
        household_selections[option] = st.checkbox(option, key=code)

    # Display selected household situations
    st.subheader('Selected Household Situations:')
    selected_households = [option for option, selected in household_selections.items() if selected]
    if selected_households:
        for selected in selected_households:
            st.write(selected)
    else:
        st.write('No options selected.')


    # LLDD, Health Problems, Other Disadvantaged Section
    st.header('LLDD, Health Problems, Other Disadvantaged')

    # Long term disability, health problem, or learning difficulties
    st.subheader('Do you consider yourself to have a long term disability, health problem or any learning difficulties? Choose the correct option. If Yes enter code in Primary LLDD or HP; you can add multiple LLDD or HP but primary must be recorded if Yes selected.')
    disability = st.radio('Choose the correct option:', ['Y', 'N'])

    # LLDD or Health Problem Types
    st.subheader('LLDD or Health Problem Type')
    columns = ['Type', 'Primary', 'Secondary', 'Tertiary']
    data = [
        ('Vision impairment (4)', 'vision_primary', 'vision_secondary', 'vision_tertiary'),
        ('Hearing impairment (5)', 'hearing_primary', 'hearing_secondary', 'hearing_tertiary'),
        ('Disability affecting mobility (6)', 'mobility_primary', 'mobility_secondary', 'mobility_tertiary'),
        ('Profound complex disabilities (7)', 'complex_primary', 'complex_secondary', 'complex_tertiary'),
        ('Social and emotional difficulties (8)', 'social_primary', 'social_secondary', 'social_tertiary'),
        ('Mental health difficulty (9)', 'mental_primary', 'mental_secondary', 'mental_tertiary'),
        ('Moderate learning difficulty (10)', 'moderate_primary', 'moderate_secondary', 'moderate_tertiary'),
        ('Severe learning difficulty (11)', 'severe_primary', 'severe_secondary', 'severe_tertiary'),
        ('Dyslexia (12)', 'dyslexia_primary', 'dyslexia_secondary', 'dyslexia_tertiary'),
        ('Dyscalculia (13)', 'dyscalculia_primary', 'dyscalculia_secondary', 'dyscalculia_tertiary'),
        ('Autism spectrum disorder (14)', 'autism_primary', 'autism_secondary', 'autism_tertiary'),
        ('Asperger\'s syndrome (15)', 'aspergers_primary', 'aspergers_secondary', 'aspergers_tertiary'),
        ('Temporary disability after illness (for example post-viral) or accident (16)', 'temporary_primary', 'temporary_secondary', 'temporary_tertiary'),
        ('Speech, Language and Communication Needs (17)', 'speech_primary', 'speech_secondary', 'speech_tertiary'),
        ('Other physical disability (18)', 'physical_primary', 'physical_secondary', 'physical_tertiary'),
        ('Other specific learning difficulty (e.g. Dyspraxia) (19)', 'specific_primary', 'specific_secondary', 'specific_tertiary'),
        ('Other medical condition (for example epilepsy, asthma, diabetes) (20)', 'medical_primary', 'medical_secondary', 'medical_tertiary'),
        ('Other learning difficulty (90)', 'other_learning_primary', 'other_learning_secondary', 'other_learning_tertiary'),
        ('Other disability (97)', 'other_disability_primary', 'other_disability_secondary', 'other_disability_tertiary'),
        ('Prefer not to say (98)', 'prefer_not_primary', 'prefer_not_secondary', 'prefer_not_tertiary')
    ]

    for label, primary, secondary, tertiary in data:
        st.write(f'**{label}**')
        st.checkbox('Primary', key=primary)
        st.checkbox('Secondary', key=secondary)
        st.checkbox('Tertiary', key=tertiary)

    # Additional information that may impact learning
    additional_info = st.text_area('Is there any other additional information that may impact on your ability to learn?')

    # Other disadvantaged sections
    st.subheader('Other disadvantaged - Ex Offender?')
    ex_offender = st.radio('', ['Y', 'N', 'Choose not to say'], key='ex_offender')

    st.subheader('Other disadvantaged - Homeless?')
    homeless = st.radio('', ['Y', 'N', 'Choose not to say'], key='homeless')

    # Referral Source Section
    st.header('Referral Source')

    # Creating columns for referral source options
    col1, col2, col3, col4 = st.columns(4)

    # Adding checkboxes for each referral source option
    with col1:
        internally_sourced = st.checkbox('Internally sourced', key='internally_sourced_1')
        recommendation = st.checkbox('Recommendation')
        event = st.checkbox('Event (please specify)')

    with col2:
        self_referral = st.checkbox('Self Referral')
        family_friends = st.checkbox('Family/ Friends')
        other = st.checkbox('Other (please specify)')

    with col3:
        website_1 = st.checkbox('Website', key='website_1')
        internally_sourced_2 = st.checkbox('Internally sourced', key='internally_sourced_2')

    with col4:
        promotional_material = st.checkbox('Promotional material')
        website_2 = st.checkbox('Website', key='website_2')

    # Text inputs for 'Event (please specify)' and 'Other (please specify)' if checked
    if event:
        event_specify = st.text_input('Please specify the event', key='event_specify')

    if other:
        other_specify = st.text_input('Please specify other source', key='other_specify')

    # Employment and Monitoring Information Section
    st.header('Employment and Monitoring Information')

    # Participant Employment Status
    st.subheader('Participant Employment Status (place an X in the applicable box)')
    employment_status = st.radio(
        "Select your employment status:",
        [
            "Unemployed (looking for work and available to start work) -> go to section A",
            "Economically Inactive (not looking for work and not available to start work) -> Go to section B",
            "Employed (including self-employed) -> go to section C"
        ]
    )

    # Section A - Unemployment details
    if "Unemployed" in employment_status:
        st.subheader('Section A - Unemployment details')
        st.text("Where a participant’s employment status is long-term unemployed proof of both unemployment and the length of unemployment must be obtained.")
        unemployment_duration = st.radio("If you are not working, how long have you been without work?", ["Up to 12 months", "12 months or longer"])
        st.write("Evidence of unemployment status (for more information look Start-Eligibility Evidence list tab)")
        unemployment_evidence = st.selectbox(
            "Select evidence type:",
            [
                "A Letter or Document from JCP or DWP",
                "A written referral from a careers service",
                "Third Party Verification or Referral form",
                "Other (please specify)"
            ]
        )
        if unemployment_evidence == "Other (please specify)":
            other_evidence = st.text_input("Please specify other evidence")

    # Section B - Economically Inactive details
    if "Economically Inactive" in employment_status:
        st.subheader('Section B - Economically Inactive details')
        inactive_status = st.radio(
            "The Participant is not employed and does not claim benefits at the time of the enrolment.",
            ["Y", "N"]
        )
        inactive_evidence_type = st.text_input("Type of evidence for Economically Inactive Status including self-declaration statement.")
        inactive_evidence_date = st.date_input("Date of issue of evidence")

    # Section C - Employment details
    if "Employed" in employment_status:
        st.subheader('Section C - Employment details')
        employer_name = st.text_input("Employer Name")
        employer_address_1 = st.text_input("Employer Address 1")
        employer_address_2 = st.text_input("Employer Address 2")
        employer_address_3 = st.text_input("Employer Address 3")
        employer_postcode = st.text_input("Employer Postcode")
        employer_contact_name = st.text_input("Main Employer Contact Name")
        employer_contact_position = st.text_input("Contact Position")
        employer_contact_email = st.text_input("Contact Email Address")
        employer_contact_phone = st.text_input("Contact Telephone Number")
        employer_edrs_number = st.text_input("Employer EDRS number")

        living_wage = st.radio("Do you earn more than the National Living Wage of £20,319.00 pa (£10.42ph for 37.5 hrs pw)?", ["Y", "N"])
        employment_hours = st.radio("Employment Hours (place an X in the applicable box)", ["0-15 hrs per week", "16+ hrs per week"])

        claiming_benefits = st.radio("Are you claiming any benefits? If so, please describe below what they are.", ["Y", "N"])
        if claiming_benefits == "Y":
            sole_claimant = st.radio("Are you the sole claimant of the benefit?", ["Y", "N"])
            benefits_list = st.multiselect(
                "Select the benefits you are claiming:",
                [
                    "Universal Credit (UC)",
                    "Job Seekers Allowance (JSA)",
                    "Employment and Support Allowance (ESA)",
                    "Incapacity Benefit (or any other sickness related benefit)",
                    "Personal Independence Payment (PIP)",
                    "Other - please state"
                ]
            )
            if "Other - please state" in benefits_list:
                other_benefit = st.text_input("Please state other benefit")
            benefit_claim_date = st.date_input("From what date has the above claim been in effect?")


    # Detailed Learning Plan Section
    st.header('Detailed Learning Plan')

    qualification_reference = st.text_input("Qualification Reference")
    region_of_work = st.text_input("Region of Work")
    qualification_course_title = st.text_input("Qualification/Course/Unit Title/Non-Regulated activity")
    awarding_body = st.text_input("Awarding Body")

    GLH = st.text_input("GLH")

    benefit_to_you = st.text_area("What is the benefit to you in completing this learning aim? Please be specific")

    planned_start_date = st.date_input("Planned Start Date")
    planned_end_date = st.date_input("Planned End Date", help="Note: Actual End Date to be recorded on 'Outcome and Progression' form at the end of the programme")
    delivery_postcode = st.text_input("Delivery Postcode")
    date_of_first_review = st.date_input("Date of first review")

    st.subheader("Progression - Indicate below the progression planned for this participant when they have completed all training")
    progression_options = st.multiselect(
        "Select progression options:",
        [
            "Progression within Work",
            "Progression into Further Education or Training",
            "Progression to Apprenticeship",
            "Progression into employment"
        ]
    )

    progression_aim = st.text_area("Please detail your progression aim")

    st.subheader("Social Outcomes - How do you rate yourself now out of 5 for the below. 5= Great 1= Poor")

    health_and_well_being = st.slider("Health and well being", 1, 5, 1)
    social_integration = st.slider("Social integration", 1, 5, 1)
    learner_self_efficacy = st.slider("Learner self-efficacy", 1, 5, 1)
    participation_in_volunteering = st.slider("Participation in volunteering", 1, 5, 1)


    # Privacy and Data Protection Information Section
    st.header('Privacy and Data Protection Information')

    # Display image
    st.image("Privacy and Data Protection Information.jpg")

    st.write("Add Y or N for any of the following boxes if you AGREE to be contacted; tick how you wish to be contacted")

    # Contact preferences
    contact_courses = st.radio("About courses/learning opportunities (fill in all boxes with either Y or N)", options=["Y", "N"])
    contact_surveys = st.radio("For surveys & research", options=["Y", "N"])
    contact_phone = st.radio("Phone", options=["Y", "N"])
    contact_email = st.radio("Email", options=["Y", "N"])
    contact_post = st.radio("Post", options=["Y", "N"])

    # Declarations Section
    st.header('Declarations')

    st.write("We hereby confirm that we have read, understood and agree with the contents of this document, and understand that the programme is funded by the Mayor of London.")

    # Participant Declaration
    st.subheader('Participant Declaration')
    st.write("I certify that I have provided all of the necessary information to confirm my eligibility for the Provision.")
    st.write("I also consent for the named Training Provider to collect further evidence, from a 3rd party Training Provider, to support a progression claim on my behalf (where applicable).")

    participant_signature = st.text_input("Participant Signature")
    participant_date = st.date_input("Date")

    participant_signature_method = st.radio(
        "I can confirm that the signature has been entered myself via the following method:",
        [
            "Wet signature of original document",
            "Inserting image of my signature (email mandate needed)",
            "Signature software such as Docusign/Adobe Sign",
            "Email declaration/mandate – this must be attached"
        ]
    )

    # Learner/Participant Eligibility & Assessment Section
    # Already done

    # Eligibility Check
    st.subheader('Eligibility Check')
    st.write("Evidence CANNOT be accepted that has been entered at a later date than Actual End Date of the start aim.")
    st.write("Evidence must be present for ALL 4 (EO1,2,3,4) of the below eligibility checks. Original documentation must have been witnessed by the Provider and preferably copies made as evidence in case of future audits.")
    st.write("For list of ALL acceptable supporting documents check 'Start-Eligibility Evidence list'")

    st.subheader('UK, EEA Nationals and Non-EEA Nationals')
    st.write("EEA Countries (as of 27/01/2021): Austria, Belgium, Bulgaria, Croatia, Republic of Cyprus, Czech Republic, Denmark, Estonia, Finland, France, Germany, Greece, Hungary, Ireland, Italy, Latvia, Lithuania, Luxembourg, Malta, Netherlands, Poland, Portugal, Romania, Slovakia, Slovenia, Spain, Sweden, Iceland, Liechtenstein, Norway.")
    st.write("Switzerland is not an EU or EEA member but is part of the single market. This means Swiss nationals have the same rights to live and work in the UK as other EEA nationals.")
    st.write("“Irish citizens in the UK hold a unique status under each country’s national law. You do not need permission to enter or remain in the UK, including a visa, any form of residence permit or employment permit”. Quote taken from [Common Travel Area Guidance](https://www.gov.uk/government/publications/common-travel-area-guidance/common-travel-area-guidance)")
    st.write("Non-EEA nationals who hold leave to enter or leave to remain with a permission to work (including status under the EUSS where they are an eligible family member of an EEA national) are also eligible for ESF support whilst in the UK.")

    # E01: Right to Live and Work in the UK
    st.subheader('E01: Right to Live and Work in the UK')
    st.subheader("UK and Irish National and European Economic Area (EEA) National?")
    nationality = st.text_input("Nationality")

    # Document Checklist
    full_uk_passport = st.checkbox("Full UK Passport")
    full_eu_passport = st.checkbox("Full EU Member Passport (must be in date - usually 10 years)")
    national_id_card = st.checkbox("National Identity Card (EU)")

    # EEA Nationals Conditions
    st.write("In order to be eligible for ESF funding, EEA Nationals must meet one of the following conditions")
    euss_settled_status = st.checkbox("a. Hold settled status granted under the EU Settlement Scheme (EUSS)")
    euss_pre_settled_status = st.checkbox("b. Hold pre-settled status granted under the European Union Settlement Scheme (EUSS)")
    points_based_system = st.checkbox("c. Hold leave to remain with permission to work granted under the new Points Based Immigration System.")

    previous_residence = st.radio("Have you been resident in the UK/EEA for the previous 3 years", options=["Y", "N"])
    country_of_birth = st.text_input("Country of Birth")
    years_in_UK= st.text_input("Years in UK")

    # Not UK, Irish or EEA National
    st.subheader('Not UK, Irish or EEA National')
    non_uk_nationality = st.text_input("Nationality ")
    passport_non_eu = st.text("Passport from non-EU member state (must be in date) AND any of the below a, b, or c")

    # Additional Documents
    indefinite_leave_letter = st.checkbox("a. Letter from the UK Immigration and Nationality Directorate granting indefinite leave to remain (settled status)")
    endorsed_passport = st.checkbox("b. Passport either endorsed 'indefinite leave to remain' – (settled status) or includes work or residency permits or visa stamps (unexpired) and all related conditions met; add details below")
    biometric_permit = st.checkbox("c. Some non-EEA nationals have an Identity Card (Biometric Permit) issued by the Home Office in place of a visa, confirming the participant’s right to stay, work or study in the UK – these cards are acceptable")

    country_of_issue = st.text_input("Country of issue")
    id_document_ref_number = st.text_input("ID Document Reference Number")
    date_of_issue = st.date_input("Date of Issue")
    date_of_expiry = st.date_input("Date of Expiry")
    additional_notes = st.text_area("Use this space for additional notes where relevant (type of Visa, restrictions, expiry etc.)")



    # E02: Proof of Age
    st.subheader('E02: Proof of Age (* all documents must be in date and if a letter is used, it must be within the last 3 months)')
    proof_of_age_options = [
        "Full Passport (EU Member State)",
        "National ID Card (EU)",
        "Firearms Certificate/ Shotgun Licence",
        "Birth/ Adoption Certificate",
        "Drivers Licence (photo card)",
        "Letter from Educational Institution* (showing DOB)",
        "Employment Contract/ Pay Slip (showing DOB)",
        "State Benefits Letter* (showing DOB)",
        "Pension Statement* (showing DOB)",
        "Northern Ireland voters card",
        "Other evidence: Please state type"
    ]
    proof_of_age = st.selectbox("Proof of Age", proof_of_age_options)

    # Conditional input for "Other evidence"
    if proof_of_age == "Other evidence: Please state type":
        other_proof_of_age = st.text_input("Please state the type of other evidence")
    date_of_issue_proof_of_age = st.date_input("Date of Issue of evidence", key="proof_of_age_date")

    # E03: Proof of Residence
    st.subheader('E03: Proof of Residence (must show the address recorded on ILP) *within the last 3 months)')
    proof_of_residence_options = [
        "Drivers Licence (photo card)",
        "Bank Statement *",
        "Pension Statement*",
        "Mortgage Statement*",
        "Utility Bill* (excluding mobile phone)",
        "Council Tax annual statement or monthly bill*",
        "Electoral Role registration evidence*",
        "Letter/confirmation from homeowner (family/lodging)",
        "Other Evidence: Please state type"
    ]
    proof_of_residence = st.selectbox("Proof of Residence", proof_of_residence_options)

    # Conditional input for "Other evidence"
    if proof_of_residence == "Other Evidence: Please state type":
        other_proof_of_residence = st.text_input("Please state the type of other evidence ")
    date_of_issue_proof_of_residence = st.date_input("Date of Issue of evidence", key="proof_of_residence_date")

    # E04: Employment Status
    st.subheader('E04: Employment Status (please select one option from below and take a copy)')

    employment_status_options = [
        "a. Latest Payslip (maximum 3 months prior to start date)",
        "b. Employment Contract",
        "c. Confirmation from the employer that the Participant is currently employed by them which must detail: Participant full name, contracted hours, start date AND date of birth or NINO",
        "d. Redundancy consultation or notice (general notice to group of staff or individual notifications) At risk of Redundancy only",
        "e. Self-employed - A submitted HMRC 'SA302' self-assessment tax declaration, with acknowledgement of receipt (within last 12 months)",
        "e. Self-employed - Records to show actual payment of Class 2 National Insurance Contributions (within last 12 months)",
        "e. Self-employed - Business records in the name of the business - evidence that a business has been established and is active / operating (within last 12 months)",
        "e. Self-employed - If registered as a Limited company: Companies House records / listed as Company Director (within last 12 months)",
        "f. Other evidence as listed in the 'Start-Eligibility Evidence list' under Employed section - State below",
        "g. Unemployed (complete the Employment section in ILP form)"
    ]

    employment_status = st.selectbox("Employment Status", employment_status_options)

    # Conditional input for "Other evidence"
    if employment_status == "f. Other evidence as listed in the 'Start-Eligibility Evidence list' under Employed section - State below":
        other_employment_status = st.text_input("Please state the type of other evidence  ")

    date_of_issue_employment_status = st.date_input("Date of Issue of evidence", key="employment_status_date")



    # Initial Assessment
    st.subheader('Initial Assessment')

    # Qualification or Training
    qualification_training = st.selectbox(
        "Are you currently undertaking a qualification or training? This could be Apprenticeship, training/qualification under ESFA/ESF funding or any other training.",
        ["Y", "N"]
    )

    if qualification_training == "Y":
        course_details = st.text_area("If yes, please supply details of the course and how this is funded")

    # Evidenced qualification levels
    st.subheader('Evidenced qualification levels:')
    self_declaration_qualification = st.selectbox(
        "Participant self declaration of highest qualification level",
        ["Below Level 1", "Level 1", "Level 2", "Level 3", "Level 4", "Level 5 and above", "No Qualifications"]
    )

    # Training Provider's declaration
    st.subheader("Training Provider's declaration:")
    provider_declaration_qualification = st.selectbox(
        "Please check the PLR and record information about prior attainment level to ensure correct recording of prior attainment, as well as ensuring no duplication of learning aims or units takes place.",
        ["Below Level 1", "Level 1", "Level 2", "Level 3", "Below Level 4", "Level 5 and above", "No Qualifications", "No Personal Learning Record"]
    )

    discrepancy_justification = st.text_area(
        "If there is a discrepancy between Participant self declaration and the PLR, please record justification for level to be reported"
    )

    # Basic Skills
    st.subheader("Does the participant have Basic Skills?")
    english_skills = st.selectbox("English", ["none", "Entry Level", "Level 1", "Level 2+"])
    maths_skills = st.selectbox("Maths", ["none", "Entry Level", "Level 1", "Level 2+"])
    esol_skills = st.selectbox("ESOL", ["none", "Entry Level", "Level 1", "Level 2+"])

    # Basic Skills Initial Assessment
    st.subheader("Basic Skills Initial Assessment:")
    maths_level = st.text_input("Initial Assessment Outcomes – record the levels achieved by the Participant - Maths Level")
    english_level = st.text_input("Initial Assessment Outcomes – record the levels achieved by the Participant - English Level")

    # Relevant Programmes
    relevant_programmes = st.selectbox(
        "Will the Participant be completing relevant Numeracy and/or Literacy programmes within their learning plan?",
        ["Y", "N"]
    )

    # Additional Learning Support
    additional_support = st.selectbox("Does the Participant require additional learning and/or learner support?", ["Y", "N"])

    if additional_support == "Y":
        support_details = st.text_area("If answered 'Yes' above, please detail how the participant will be supported")


    # Current Skills, Experience and IAG
    st.subheader('Current Skills, Experience and IAG')

    # Highest Level of Education at start
    education_level = st.selectbox(
        "Highest Level of Education at start",
        [
            "ISCED 0 - Lacking Foundation skills (below Primary Education)",
            "ISCED 1 - Primary Education",
            "ISCED 2 - GCSE D-G or 3-1/BTEC Level 1/Functional Skills Level 1",
            "ISCED 3 - GCSE A-C or 9-4/AS or A Level/NVQ or BTEC Level 2 or 3",
            "ISCED 4 - N/A",
            "ISCED 5 to 8 - BTEC Level 5 or NVQ Level 4, Foundation Degree, BA, MA or equivalent"
        ]
    )

    # Other Information
    st.subheader("Other Information:")

    # Current Job Role and Activities
    current_job_role = st.text_area("What is your current job role and what are your day to day activities?")

    # Career Aspirations
    career_aspirations = st.text_area("What are your career aspirations?")

    # Training/Qualifications Needed
    training_needed = st.text_area("What training/qualifications do you need to progress further in your career? (Planned and future training)")

    # Barriers to Career Aspirations
    career_barriers = st.text_area("What are the barriers to achieving your career aspirations and goals?")

    # Courses/Programmes/Activity Available
    courses_available = st.text_area("What courses/programmes/activity are available to the Participant in order to meet their and their employers' needs?")

    # Induction Checklist
    st.subheader('Induction Checklist')
    st.write("The following have been discussed with the Participant as part of the induction onto the programme")

    induction_checklist_options = [
        "This programme is funded by the ESFA",
        "Describe the programme content and delivery expectation",
        "Equality and Diversity Policy/ Procedure and point of contact",
        "Health and Safety Policy/ Procedure and point of contact",
        "Safeguarding Policy/ Procedure and point of contact",
        "PREVENT and point of contact (including British Values)",
        "Disciplinary, Appeal and Grievance Policy/ Procedures",
        "Plagiarism, Cheating Policy/ Procedure",
        "Terms and Conditions of Learning"
    ]

    induction_checklist = {}
    for option in induction_checklist_options:
        induction_checklist[option] = st.checkbox(option)

    # # Display the user's selections (optional)
    # st.write("Induction Checklist Selections:")
    # for option, selected in induction_checklist.items():
    #     if selected:
    #         st.write(f"- {option}")


    st.header('Declarations')
    st.text(
        'We hereby confirm that we have read, understood and agree with the contents of this document, and understand that the programme is funded by the Mayor of London.'
    )


    st.subheader('Participant Declaration')
    participant_declaration = st.text_area(
        'Participant Declaration',
        'I certify that I have provided all of the necessary information to confirm my eligibility for the Funded Provision.'
    )

    st.subheader('Participant Signature')
    signature_method = st.radio('Select how the signature has been entered:', [
        'Wet signature of original document',
        'Inserting image of my signature (email mandate needed)',
        'Signature software such as Docusign/Adobe Sign',
        'Email declaration/mandate – this must be attached'
    ])
    st.text("Signature:")
    participant_signature = st_canvas(
        fill_color="rgba(255, 255, 255, 1)",  
        stroke_width=5,
        stroke_color="rgb(0, 0, 0)",  # Black stroke color
        background_color="white",  # White background color
        width=400,
        height=150,
        drawing_mode="freedraw",
        key="participant_signature_canvas",
    )

    date_signed = st.date_input(
    label="Sign Date",
    value=datetime(2000, 1, 1),  # Default date
    min_value=date(1900, 1, 1),  # Minimum selectable date
    max_value=date(2025, 12, 31),  # Maximum selectable date
    help="Choose a date"  # Tooltip text
)
    
# #############################################################
# Submit
    submit_button = st.button('Submit')
    if submit_button:
        placeholder_values = {
            'p1': partnership,
            'p2': learner_name,
            'p3': qualification,
        }

        # Define input and output paths
        template_file = "ph esfa.docx"
        modified_file = "Filled_EFSA_form.docx"

        if participant_signature.image_data is not None:
            # Convert the drawing to a PIL image and save it
            signature_path = 'signature_image.png'
            signature_image = PILImage.fromarray(participant_signature.image_data.astype('uint8'), 'RGBA')
            signature_image.save(signature_path)
            # st.success("Signature image saved!")

            replace_placeholders(template_file, modified_file, placeholder_values, signature_path)
            # st.success(f"Template modified and saved as {modified_file}")
        else:
            st.warning("Please draw your signature.")
        st.success('Form submitted successfully!')

def resize_image_to_fit_cell(image_path, max_width, max_height):
    with PILImage.open(image_path) as img:
        img.thumbnail((max_width, max_height), PILImage.Resampling.LANCZOS)
        return img

def replace_placeholders(template_file, modified_file, placeholder_values, signature_path):
    # Copy the template file to a new file
    shutil.copy(template_file, modified_file)

    doc = Document(modified_file)

    for para in doc.paragraphs:
        for placeholder, value in placeholder_values.items():
            if re.search(r'\b' + re.escape(placeholder) + r'\b', para.text):
                para.text = re.sub(r'\b' + re.escape(placeholder) + r'\b', value, para.text)

    # Handling tables, if placeholders are present in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for placeholder, value in placeholder_values.items():
                        if re.search(r'\b' + re.escape(placeholder) + r'\b', para.text):
                            para.text = re.sub(r'\b' + re.escape(placeholder) + r'\b', value, para.text)

    # Handle signature placeholder separately
    for para in doc.paragraphs:
        if 'SIGNATURE_PLACEHOLDER' in para.text:
            para.text = para.text.replace('SIGNATURE_PLACEHOLDER', '')
            resized_image = resize_image_to_fit_cell(signature_path, 200, 55)
            resized_image_path = 'resized_signature_image.png'
            resized_image.save(resized_image_path)
            para.add_run().add_picture(resized_image_path, width=Inches(2))

    doc.save(modified_file)

    # file download button
    with open(modified_file, 'rb') as f:
        file_contents = f.read()
        st.download_button(
            label="Download File",
            data=file_contents,
            file_name=modified_file,
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

if __name__ == '__main__':
    app()
